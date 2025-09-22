"""Document creation utilities for generating Word documents."""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from .text_cleaner import TextCleaner


class DocumentCreator:
    """Utility class for creating Word documents from OCR text."""

    def __init__(self, text_direction='rtl', encoding='utf-8'):
        """
        Initialize the document creator.

        Args:
            text_direction: Text direction ('ltr' or 'rtl')
            encoding: Text encoding
        """
        self.text_direction = text_direction
        self.encoding = encoding
        self.text_cleaner = TextCleaner()

    def create_word_document_with_pages(self, page_texts, output_path, page_numbers=None):
        """
        Create a Word document from a list of page texts, maintaining 1:1 page mapping

        Args:
            page_texts: List of text strings, one for each page
            output_path: Path to save the Word document
            page_numbers: Optional list of original page numbers from OCR
        """
        doc = Document()

        # Set up document formatting for Arabic/RTL text
        style = doc.styles['Normal']
        font = style.font

        # Use Arabic-friendly fonts
        if self.text_direction == 'rtl':
            font.name = 'Arial Unicode MS'  # Better Arabic support
        else:
            font.name = 'Arial'
        font.size = Pt(12)  # Slightly larger for Arabic readability

        # Add title page
        title = doc.add_heading('OCR Extracted Text', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add document info
        info_para = doc.add_paragraph(f'Source: {Path(output_path).stem}')
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        info_para2 = doc.add_paragraph(f'Total pages: {len(page_texts)}')
        info_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        total_chars = sum(len(page) for page in page_texts)
        info_para3 = doc.add_paragraph(f'Total characters: {total_chars:,}')
        info_para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        info_para4 = doc.add_paragraph(f'Text direction: {self.text_direction.upper()}')
        info_para4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()

        # Add each page
        for idx, page_text in enumerate(page_texts):
            # Use provided page numbers or sequential numbering
            display_page_num = page_numbers[idx] if page_numbers and idx < len(page_numbers) else idx + 1

            # Add page header
            header = doc.add_paragraph(f'Page {display_page_num}')
            if self.text_direction == 'rtl':
                header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            header.runs[0].font.size = Pt(10)
            header.runs[0].font.italic = True

            # Clean and add page text
            cleaned_text = self.text_cleaner.clean_text(page_text)

            # Split into paragraphs and add
            paragraphs = cleaned_text.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    p = doc.add_paragraph(para)
                    p.paragraph_format.space_after = Pt(6)

                    # Set text direction
                    if self.text_direction == 'rtl':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        # Add RTL paragraph properties
                        pPr = p._element.get_or_add_pPr()
                        bidi = parse_xml(r'<w:bidi %s/>' % nsdecls('w'))
                        pPr.append(bidi)
                    else:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add page break except for last page
            if idx < len(page_texts) - 1:
                doc.add_page_break()

            # Show progress for large documents
            if (idx + 1) % 10 == 0:
                print(f"Processing page {idx + 1}/{len(page_texts)}...")

        # Save document
        doc.save(output_path)
        print(f"Word document saved to: {output_path}")

    def create_word_document(self, text, output_path, chars_per_page=3000):
        """
        Create a Word document from text, splitting into pages

        Args:
            text: The text to write to the document
            output_path: Path to save the Word document
            chars_per_page: Approximate characters per page (default 3000)
        """
        doc = Document()

        # Set up document formatting
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Add title page
        title = doc.add_heading('OCR Extracted Text', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f'Source: {Path(output_path).stem}')
        doc.add_paragraph(f'Total characters: {len(text):,}')
        doc.add_paragraph(f'Approximate pages: {(len(text) // chars_per_page) + 1}')
        doc.add_page_break()

        # Split text into chunks for pages
        paragraphs = text.split('\n\n')
        current_page_text = []
        current_char_count = 0
        page_number = 1

        for i, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # Check if adding this paragraph would exceed page limit
            if current_char_count + len(paragraph) > chars_per_page and current_page_text:
                # Add page header
                header = doc.add_paragraph(f'Page {page_number}')
                header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                header.runs[0].font.size = Pt(10)
                header.runs[0].font.italic = True

                # Add current page content
                for para in current_page_text:
                    p = doc.add_paragraph(para)
                    # Add some spacing between paragraphs
                    p.paragraph_format.space_after = Pt(6)

                # Add page break
                doc.add_page_break()
                page_number += 1

                # Reset for next page
                current_page_text = [paragraph]
                current_char_count = len(paragraph)
            else:
                current_page_text.append(paragraph)
                current_char_count += len(paragraph)

            # Show progress for large documents
            if i % 100 == 0 and i > 0:
                print(f"Processing paragraph {i}/{len(paragraphs)}...")

        # Add remaining content
        if current_page_text:
            # Add final page header
            header = doc.add_paragraph(f'Page {page_number}')
            header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            header.runs[0].font.size = Pt(10)
            header.runs[0].font.italic = True

            for para in current_page_text:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(6)

        # Save document
        doc.save(output_path)
        print(f"Word document saved to: {output_path}")