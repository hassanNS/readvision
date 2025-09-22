from google.cloud import translate_v2 as translate

client = translate.Client.from_service_account_json("gcp.json")

#!/usr/bin/env python3
"""
Google Cloud Vision API OCR Script for Large PDFs
Processes multi-page PDFs using batch operations for efficiency
"""

import os
import json
import time
import re
import argparse
import sys
from pathlib import Path
from google.cloud import vision
from google.cloud import storage
from google.protobuf import json_format
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class PDFOCRProcessor:
    def __init__(self, credentials_path=None, bucket_name=None):
        """
        Initialize the OCR processor

        Args:
            credentials_path: Path to Google Cloud credentials JSON file
            bucket_name: Google Cloud Storage bucket name for temporary storage
        """
        if credentials_path:
            os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = credentials_path

        self.vision_client = vision.ImageAnnotatorClient()
        self.storage_client = storage.Client()

        # Create or use existing bucket for batch operations
        if bucket_name:
            self.bucket_name = bucket_name
        else:
            self.bucket_name = f"ocr-temp-{int(time.time())}"
            self._create_bucket()

    def _create_bucket(self):
        """Create a temporary bucket for batch operations"""
        try:
            bucket = self.storage_client.create_bucket(self.bucket_name)
            print(f"Created bucket: {self.bucket_name}")
        except Exception as e:
            print(f"Bucket creation failed: {e}")
            # Use existing bucket or handle error

    def clean_text(self, text):
        """
        Clean OCR text output

        Args:
            text: Raw OCR text

        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)

        # Fix common OCR errors
        text = text.replace('|', 'I')  # Common OCR mistake
        text = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', text)  # Add space between camelCase

        # Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')

        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;!?])', r'\1', text)
        text = re.sub(r'([.,;!?])(?=[A-Za-z])', r'\1 ', text)

        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()

    def debug_page_order(self, page_data):
        """
        Debug function to show page ordering information

        Args:
            page_data: List of (page_number, text) tuples
        """
        print(f"\n🔍 Debug: Page Order Information")
        print(f"Total pages found: {len(page_data)}")

        if page_data:
            page_numbers = [page_num for page_num, text in page_data]
            print(f"Page numbers found: {sorted(set(page_numbers))}")
            print(f"Min page: {min(page_numbers)}, Max page: {max(page_numbers)}")

            # Check for duplicates
            duplicates = [num for num in set(page_numbers) if page_numbers.count(num) > 1]
            if duplicates:
                print(f"⚠️  Duplicate page numbers found: {duplicates}")

            # Check for missing pages
            expected_range = set(range(min(page_numbers), max(page_numbers) + 1))
            actual_pages = set(page_numbers)
            missing = expected_range - actual_pages
            if missing:
                print(f"⚠️  Missing page numbers: {sorted(missing)}")

            # Show first few pages for verification
            sorted_data = sorted(page_data, key=lambda x: x[0])
            print(f"\nFirst 3 pages preview:")
            for i, (page_num, text) in enumerate(sorted_data[:3]):
                preview = text.strip()[:100].replace('\n', ' ')
                print(f"  Page {page_num}: {preview}...")
        print("-" * 50)

    def create_word_document_with_pages(self, page_texts, output_path, page_numbers=None):
        """
        Create a Word document from a list of page texts, maintaining 1:1 page mapping

        Args:
            page_texts: List of text strings, one for each page
            output_path: Path to save the Word document
            page_numbers: Optional list of original page numbers from OCR
        """
        doc = Document()

        # Get text direction and encoding settings
        text_direction = getattr(self, 'text_direction', 'rtl')
        encoding = getattr(self, 'encoding', 'utf-8')

        # Set up document formatting for Arabic/RTL text
        style = doc.styles['Normal']
        font = style.font

        # Use Arabic-friendly fonts
        if text_direction == 'rtl':
            font.name = 'Arial Unicode MS'  # Better Arabic support
        else:
            font.name = 'Arial'
        font.size = Pt(12)  # Slightly larger for Arabic readability

        # Add title page
        title = doc.add_heading('OCR Extracted Text', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add document info
        info_para = doc.add_paragraph(f'Source: {Path(output_path).stem}')
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        info_para2 = doc.add_paragraph(f'Total pages: {len(page_texts)}')
        info_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        total_chars = sum(len(page) for page in page_texts)
        info_para3 = doc.add_paragraph(f'Total characters: {total_chars:,}')
        info_para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        info_para4 = doc.add_paragraph(f'Text direction: {text_direction.upper()}')
        info_para4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()

        # Add each page
        for idx, page_text in enumerate(page_texts):
            # Use provided page numbers or sequential numbering
            display_page_num = page_numbers[idx] if page_numbers and idx < len(page_numbers) else idx + 1

            # Add page header
            header = doc.add_paragraph(f'Page {display_page_num}')
            if text_direction == 'rtl':
                header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            header.runs[0].font.size = Pt(10)
            header.runs[0].font.italic = True

            # Clean and add page text
            cleaned_text = self.clean_text(page_text)

            # Split into paragraphs and add
            paragraphs = cleaned_text.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    p = doc.add_paragraph(para)
                    p.paragraph_format.space_after = Pt(6)

                    # Set text direction
                    if text_direction == 'rtl':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        # Add RTL paragraph properties
                        pPr = p._element.get_or_add_pPr()
                        bidi = parse_xml(r'<w:bidi %s/>' % nsdecls('w'))
                        pPr.append(bidi)
                    else:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add page break except for last page
            if idx < len(page_texts) - 1:
                doc.add_page_break()

            # Show progress for large documents
            if (idx + 1) % 10 == 0:
                print(f"Processing page {idx + 1}/{len(page_texts)}...")

        # Save document
        doc.save(output_path)
        print(f"Word document saved to: {output_path}")

    def create_word_document(self, text, output_path, chars_per_page=3000):
        """
        Create a Word document from text, splitting into pages

        Args:
            text: The text to write to the document
            output_path: Path to save the Word document
            chars_per_page: Approximate characters per page (default 3000)
        """
        doc = Document()

        # Set up document formatting
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Add title page
        title = doc.add_heading('OCR Extracted Text', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f'Source: {Path(output_path).stem}')
        doc.add_paragraph(f'Total characters: {len(text):,}')
        doc.add_paragraph(f'Approximate pages: {(len(text) // chars_per_page) + 1}')
        doc.add_page_break()

        # Split text into chunks for pages
        paragraphs = text.split('\n\n')
        current_page_text = []
        current_char_count = 0
        page_number = 1

        for i, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # Check if adding this paragraph would exceed page limit
            if current_char_count + len(paragraph) > chars_per_page and current_page_text:
                # Add page header
                header = doc.add_paragraph(f'Page {page_number}')
                header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                header.runs[0].font.size = Pt(10)
                header.runs[0].font.italic = True

                # Add current page content
                for para in current_page_text:
                    p = doc.add_paragraph(para)
                    # Add some spacing between paragraphs
                    p.paragraph_format.space_after = Pt(6)

                # Add page break
                doc.add_page_break()
                page_number += 1

                # Reset for next page
                current_page_text = [paragraph]
                current_char_count = len(paragraph)
            else:
                current_page_text.append(paragraph)
                current_char_count += len(paragraph)

            # Show progress for large documents
            if i % 100 == 0 and i > 0:
                print(f"Processing paragraph {i}/{len(paragraphs)}...")

        # Add remaining content
        if current_page_text:
            # Add final page header
            header = doc.add_paragraph(f'Page {page_number}')
            header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            header.runs[0].font.size = Pt(10)
            header.runs[0].font.italic = True

            for para in current_page_text:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(6)

        # Save document
        doc.save(output_path)
        print(f"Word document saved to: {output_path}")

    def process_small_pdf(self, pdf_path, output_path):
        """
        Process PDFs with fewer than 5 pages using synchronous requests

        Args:
            pdf_path: Path to input PDF
            output_path: Path to output text file
        """
        print(f"Processing small PDF: {pdf_path}")

        page_texts = []

        # Read PDF file
        with open(pdf_path, 'rb') as file:
            content = file.read()

        # Configure OCR parameters with language hint
        feature = vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)

        # Get language hint
        language_hint = getattr(self, 'language_hint', 'ar')

        # Create image context with language hints
        image_context = vision.ImageContext(language_hints=[language_hint])

        # Create request
        request = vision.AnnotateFileRequest(
            input_config=vision.InputConfig(
                content=content,
                mime_type='application/pdf'
            ),
            features=[feature],
            image_context=image_context
        )

        # Perform OCR
        response = self.vision_client.batch_annotate_files(requests=[request])

        # Extract text from all pages with page numbers for correct ordering
        page_data = []  # Store (page_number, text) tuples

        for page_response in response.responses[0].responses:
            if page_response.full_text_annotation:
                # Try to get page number from context
                page_number = 1  # default
                if hasattr(page_response, 'context') and page_response.context:
                    page_number = getattr(page_response.context, 'pageNumber', len(page_data) + 1)
                else:
                    page_number = len(page_data) + 1

                page_text = page_response.full_text_annotation.text
                page_data.append((page_number, page_text))

        # Sort pages by page number to ensure correct order
        page_data.sort(key=lambda x: x[0])

        # Debug page ordering if enabled
        if getattr(self, 'debug', False):
            self.debug_page_order(page_data)

        # Extract text and page numbers in correct order
        page_texts = [text for page_num, text in page_data]
        page_numbers = [page_num for page_num, text in page_data]

        print(f"Processed {len(page_texts)} pages in correct order")
        if page_numbers:
            print(f"Page number range: {min(page_numbers)} to {max(page_numbers)}")

        # Create Word document with page-by-page mapping
        word_output_path = output_path.replace('.txt', '.docx')
        self.create_word_document_with_pages(page_texts, word_output_path, page_numbers)

        # Also save as combined text file if needed
        all_text = []
        for page_text in page_texts:
            all_text.append(self.clean_text(page_text))

        combined_text = '\n\n--- PAGE BREAK ---\n\n'.join(all_text)

        # Use specified encoding
        encoding = getattr(self, 'encoding', 'utf-8')
        with open(output_path, 'w', encoding=encoding) as f:
            f.write(combined_text)

        print(f"OCR completed. Output saved to: {word_output_path} and {output_path}")

    def process_large_pdf(self, pdf_path, output_path):
        """
        Process large PDFs using asynchronous batch operations

        Args:
            pdf_path: Path to input PDF
            output_path: Path to output text file
        """
        print(f"Processing large PDF: {pdf_path}")

        # Upload PDF to GCS
        blob_name = f"input/{Path(pdf_path).name}"
        bucket = self.storage_client.bucket(self.bucket_name)
        blob = bucket.blob(blob_name)
        blob.upload_from_filename(pdf_path)

        gcs_source_uri = f"gs://{self.bucket_name}/{blob_name}"
        gcs_destination_uri = f"gs://{self.bucket_name}/output/"

        # Configure batch request
        gcs_source = vision.GcsSource(uri=gcs_source_uri)
        input_config = vision.InputConfig(
            gcs_source=gcs_source,
            mime_type='application/pdf'
        )

        gcs_destination = vision.GcsDestination(uri=gcs_destination_uri)
        output_config = vision.OutputConfig(
            gcs_destination=gcs_destination,
            batch_size=100  # Max pages per output file
        )

        # Get language hint and create image context
        language_hint = getattr(self, 'language_hint', 'ar')
        image_context = vision.ImageContext(language_hints=[language_hint])

        async_request = vision.AsyncAnnotateFileRequest(
            features=[vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)],
            input_config=input_config,
            output_config=output_config,
            image_context=image_context
        )

        # Start async batch operation
        operation = self.vision_client.async_batch_annotate_files(
            requests=[async_request]
        )

        print("Waiting for operation to complete...")
        operation.result(timeout=600)  # 10 minute timeout

        # Download and process results
        page_data = []  # Store (page_number, text) tuples

        # List all output files and sort them to maintain order
        blobs = list(bucket.list_blobs(prefix='output/'))
        blobs.sort(key=lambda x: x.name)

        for blob in blobs:
            if blob.name.endswith('.json'):
                # Download JSON result
                json_content = blob.download_as_text()
                response = json.loads(json_content)

                # Extract text from each page with page number
                for page_response in response['responses']:
                    if 'fullTextAnnotation' in page_response and 'context' in page_response:
                        page_number = page_response['context'].get('pageNumber', 0)
                        page_text = page_response['fullTextAnnotation']['text']
                        page_data.append((page_number, page_text))
                    elif 'fullTextAnnotation' in page_response:
                        # Fallback if no context/pageNumber available
                        page_text = page_response['fullTextAnnotation']['text']
                        page_data.append((len(page_data) + 1, page_text))

        # Sort pages by page number to ensure correct order
        page_data.sort(key=lambda x: x[0])

        # Debug page ordering if enabled
        if getattr(self, 'debug', False):
            self.debug_page_order(page_data)

        # Extract text and page numbers in correct order
        page_texts = [text for page_num, text in page_data]
        page_numbers = [page_num for page_num, text in page_data]

        print(f"Processed {len(page_texts)} pages in correct order")
        if page_numbers:
            print(f"Page number range: {min(page_numbers)} to {max(page_numbers)}")

        # Create Word document with page-by-page mapping
        word_output_path = output_path.replace('.txt', '.docx')
        self.create_word_document_with_pages(page_texts, word_output_path, page_numbers)

        # Also save as combined text file if needed
        all_text = []
        for page_text in page_texts:
            all_text.append(self.clean_text(page_text))

        combined_text = '\n\n--- PAGE BREAK ---\n\n'.join(all_text)

        # Use specified encoding
        encoding = getattr(self, 'encoding', 'utf-8')
        with open(output_path, 'w', encoding=encoding) as f:
            f.write(combined_text)

        # Cleanup GCS
        # self._cleanup_gcs()

        print(f"OCR completed. Output saved to: {word_output_path} and {output_path}")

    def _cleanup_gcs(self):
        """Clean up temporary files from GCS bucket"""
        bucket = self.storage_client.bucket(self.bucket_name)
        blobs = bucket.list_blobs()
        for blob in blobs:
            blob.delete()

    def get_pdf_page_count(self, pdf_path):
        """Get number of pages in PDF"""
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return len(reader.pages)

    def process_pdf(self, pdf_path, output_path='output.txt', chars_per_page=3000,
                    text_direction='rtl', encoding='utf-8', language_hint='ar', debug=False):
        """
        Main method to process a PDF file

        Args:
            pdf_path: Path to input PDF
            output_path: Path to output text file
            chars_per_page: Characters per page in Word document
            text_direction: Text direction ('ltr' or 'rtl')
            encoding: Text file encoding
            language_hint: Language hint for OCR processing
            debug: Enable debug output for page ordering
        """
        # Store parameters for later use
        self.chars_per_page = chars_per_page
        self.text_direction = text_direction
        self.encoding = encoding
        self.language_hint = language_hint
        self.debug = debug

        # Check if file exists
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")

        # Get page count
        page_count = self.get_pdf_page_count(pdf_path)
        print(f"PDF has {page_count} pages")

        # Choose processing method based on size
        if page_count <= 5:
            self.process_small_pdf(pdf_path, output_path)
        else:
            self.process_large_pdf(pdf_path, output_path)


def main():
    """Main function with command line argument support"""

    # Set up command line argument parsing
    parser = argparse.ArgumentParser(
        description='OCR PDF processor that creates Word documents with page-by-page mapping',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python translate.py assets/document.pdf output.txt
  python translate.py assets/document.pdf output.txt --credentials gcp.json
  python translate.py assets/document.pdf output.txt --bucket my-bucket
  python translate.py assets/arabic.pdf output.txt --text-direction rtl --encoding utf-8
  python translate.py assets/english.pdf output.txt --text-direction ltr --language-hint en
        '''
    )

    parser.add_argument('pdf_path',
                       help='Path to the input PDF file')
    parser.add_argument('output_path',
                       help='Path for the output text file (Word doc will be created automatically)')
    parser.add_argument('--credentials', '-c',
                       default='gcp.json',
                       help='Path to Google Cloud credentials JSON file (default: gcp.json)')
    parser.add_argument('--bucket', '-b',
                       help='Google Cloud Storage bucket name for temporary storage (optional)')
    parser.add_argument('--chars-per-page',
                       type=int,
                       default=3000,
                       help='Characters per page for legacy text-based splitting (default: 3000)')
    parser.add_argument('--text-direction', '-d',
                       choices=['ltr', 'rtl'],
                       default='rtl',
                       help='Text direction: ltr (left-to-right) or rtl (right-to-left) (default: rtl for Arabic)')
    parser.add_argument('--encoding', '-e',
                       default='utf-8',
                       help='Text file encoding (default: utf-8 for Arabic support)')
    parser.add_argument('--language-hint',
                       default='ar',
                       help='Language hint for OCR processing (default: ar for Arabic)')
    parser.add_argument('--debug',
                       action='store_true',
                       help='Enable debug output for page ordering')

    args = parser.parse_args()

    # Validate input file exists
    if not os.path.exists(args.pdf_path):
        print(f"Error: PDF file not found: {args.pdf_path}")
        sys.exit(1)

    # Validate credentials file exists
    if not os.path.exists(args.credentials):
        print(f"Error: Credentials file not found: {args.credentials}")
        sys.exit(1)

    # Configuration from command line arguments
    CREDENTIALS_PATH = args.credentials
    BUCKET_NAME = args.bucket
    PDF_PATH = args.pdf_path
    OUTPUT_PATH = args.output_path
    CHARS_PER_PAGE = args.chars_per_page
    TEXT_DIRECTION = args.text_direction
    ENCODING = args.encoding
    LANGUAGE_HINT = args.language_hint
    DEBUG = args.debug

    print(f"Processing PDF: {PDF_PATH}")
    print(f"Output will be saved to: {OUTPUT_PATH}")
    print(f"Word document will be saved to: {OUTPUT_PATH.replace('.txt', '.docx')}")
    print(f"Text direction: {TEXT_DIRECTION.upper()}")
    print(f"Encoding: {ENCODING}")
    print(f"Language hint: {LANGUAGE_HINT}")
    if DEBUG:
        print(f"🔍 Debug mode: ENABLED")

    # Initialize processor
    processor = PDFOCRProcessor(
        credentials_path=CREDENTIALS_PATH,
        bucket_name=BUCKET_NAME
    )

    try:
        # Process PDF
        processor.process_pdf(
            PDF_PATH,
            OUTPUT_PATH,
            chars_per_page=CHARS_PER_PAGE,
            text_direction=TEXT_DIRECTION,
            encoding=ENCODING,
            language_hint=LANGUAGE_HINT,
            debug=DEBUG
        )

        # Read and display sample of output
        with open(OUTPUT_PATH, 'r', encoding=ENCODING) as f:
            sample = f.read(500)
            print("\nSample of extracted text:")
            print("-" * 50)
            print(sample)
            print("-" * 50)

        # Inform about Word document
        word_path = OUTPUT_PATH.replace('.txt', '.docx')
        print(f"\nWord document created: {word_path}")
        print("Each page of the PDF maps to a page in the Word document.")

    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()

